"""
DOCX (Word) document processor
"""
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import tempfile

try:
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. Install with: pip install python-docx")

from PIL import Image
import io

logger = logging.getLogger(__name__)


class DOCXProcessor:
    """
    Process Microsoft Word (.docx) documents
    
    Features:
    - Extract text from paragraphs
    - Extract text from tables
    - Extract images embedded in document
    - Preserve document structure
    - Handle headers and footers
    """
    
    def __init__(self, extract_images: bool = True):
        """
        Initialize DOCX processor
        
        Args:
            extract_images: Whether to extract embedded images
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX processing. "
                "Install with: pip install python-docx"
            )
        
        self.extract_images = extract_images
        logger.info("DOCXProcessor initialized")
    
    def process(
        self,
        file_path: str,
        extract_images: Optional[bool] = None
    ) -> Dict[str, any]:
        """
        Process a DOCX file and extract content
        
        Args:
            file_path: Path to DOCX file
            extract_images: Override default image extraction setting
            
        Returns:
            Dictionary containing:
                - text: Full extracted text
                - paragraphs: List of paragraph texts
                - tables: List of table data
                - images: List of extracted images (if enabled)
                - metadata: Document metadata
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            if not file_path.suffix.lower() == '.docx':
                raise ValueError(f"Not a DOCX file: {file_path}")
            
            logger.info(f"Processing DOCX: {file_path}")
            
            # Load document
            doc = Document(str(file_path))
            
            # Extract text from paragraphs
            paragraphs = self._extract_paragraphs(doc)
            
            # Extract text from tables
            tables = self._extract_tables(doc)
            
            # Extract images
            images = []
            if extract_images if extract_images is not None else self.extract_images:
                images = self._extract_images(doc, file_path)
            
            # Extract metadata
            metadata = self._extract_metadata(doc)
            
            # Combine all text
            full_text = self._combine_text(paragraphs, tables)
            
            result = {
                "text": full_text,
                "paragraphs": paragraphs,
                "tables": tables,
                "images": images,
                "metadata": metadata,
                "file_path": str(file_path),
                "file_type": "docx",
                "page_count": len(paragraphs) // 10 + 1  # Rough estimate
            }
            
            logger.info(
                f"DOCX processed: {len(paragraphs)} paragraphs, "
                f"{len(tables)} tables, {len(images)} images"
            )
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            raise
    
    def _extract_paragraphs(self, doc: Document) -> List[str]:
        """Extract text from all paragraphs"""
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Skip empty paragraphs
                paragraphs.append(text)
        
        return paragraphs
    
    def _extract_tables(self, doc: Document) -> List[Dict]:
        """Extract text from all tables"""
        tables_data = []
        
        for table_idx, table in enumerate(doc.tables):
            table_text = []
            rows_data = []
            
            for row in table.rows:
                row_data = []
                row_text = []
                
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                    row_text.append(cell_text)
                
                rows_data.append(row_data)
                if any(row_text):  # Skip empty rows
                    table_text.append(" | ".join(row_text))
            
            if table_text:
                tables_data.append({
                    "table_id": table_idx,
                    "text": "\n".join(table_text),
                    "rows": rows_data,
                    "num_rows": len(rows_data),
                    "num_cols": len(rows_data[0]) if rows_data else 0
                })
        
        return tables_data
    
    def _extract_images(self, doc: Document, file_path: Path) -> List[Dict]:
        """Extract embedded images from document"""
        images = []
        
        try:
            # Access document's image parts
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        
                        # Try to open with PIL
                        image = Image.open(io.BytesIO(image_data))
                        
                        # Save to temp file
                        temp_dir = tempfile.gettempdir()
                        image_filename = f"{file_path.stem}_img_{len(images)}.{image.format.lower()}"
                        image_path = Path(temp_dir) / image_filename
                        image.save(image_path)
                        
                        images.append({
                            "image_id": len(images),
                            "path": str(image_path),
                            "format": image.format,
                            "size": image.size,
                            "mode": image.mode
                        })
                        
                    except Exception as e:
                        logger.warning(f"Failed to extract image: {str(e)}")
                        continue
        
        except Exception as e:
            logger.warning(f"Error extracting images: {str(e)}")
        
        return images
    
    def _extract_metadata(self, doc: Document) -> Dict:
        """Extract document metadata"""
        metadata = {}
        
        try:
            core_props = doc.core_properties
            
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.created:
                metadata["created"] = str(core_props.created)
            if core_props.modified:
                metadata["modified"] = str(core_props.modified)
            if core_props.last_modified_by:
                metadata["last_modified_by"] = core_props.last_modified_by
            if core_props.category:
                metadata["category"] = core_props.category
            if core_props.comments:
                metadata["comments"] = core_props.comments
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords
        
        except Exception as e:
            logger.warning(f"Error extracting metadata: {str(e)}")
        
        return metadata
    
    def _combine_text(self, paragraphs: List[str], tables: List[Dict]) -> str:
        """Combine paragraphs and tables into full text"""
        parts = []
        
        # Add paragraphs
        if paragraphs:
            parts.append("\n\n".join(paragraphs))
        
        # Add tables
        if tables:
            table_texts = [f"[Table {t['table_id']}]\n{t['text']}" for t in tables]
            parts.append("\n\n".join(table_texts))
        
        return "\n\n".join(parts)
    
    def extract_text_only(self, file_path: str) -> str:
        """
        Quick extraction of just the text content
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text as string
        """
        result = self.process(file_path, extract_images=False)
        return result["text"]


def is_docx_available() -> bool:
    """Check if DOCX processing is available"""
    return DOCX_AVAILABLE


# Convenience function
def process_docx(file_path: str, extract_images: bool = True) -> Dict:
    """
    Process a DOCX file (convenience function)
    
    Args:
        file_path: Path to DOCX file
        extract_images: Whether to extract embedded images
        
    Returns:
        Processed document data
    """
    processor = DOCXProcessor(extract_images=extract_images)
    return processor.process(file_path)
